#!/usr/bin/env python3
"""Generate post-transfer company credentials Word document (fill-in template)."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date

OUTPUT = r"e:\vigilance git\Vigilance\docs\Vigilance_Company_Account_Details.docx"


def add_table_rows(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def build():
    doc = Document()
    title = doc.add_heading("Vigilance VMS — Company Account Details", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Document date: {date.today().strftime('%B %d, %Y')}")
    doc.add_paragraph(
        "CONFIDENTIAL — Store securely. Do not commit to git or share publicly. "
        "Fill in each field after completing the transfer."
    )

    doc.add_heading("Primary company contact", 1)
    add_table_rows(doc, ["Field", "Value"], [
        ("Organization", "Vigilance CFCICI"),
        ("Admin email", "vigilancecfcici@gmail.com"),
        ("Password", "[STORE IN PASSWORD MANAGER — NOT IN THIS FILE]"),
        ("Notes", ""),
    ])

    doc.add_heading("GitHub", 1)
    add_table_rows(doc, ["Field", "Value"], [
        ("Account URL", "https://github.com/"),
        ("Organization / username", ""),
        ("Repository URL", ""),
        ("Default branch", "main"),
        ("Actions secrets configured", "YES / NO"),
    ])

    doc.add_heading("Supabase", 1)
    add_table_rows(doc, ["Field", "Value"], [
        ("Dashboard", "https://supabase.com/dashboard"),
        ("Organization", ""),
        ("Project name", ""),
        ("Project ref (ID)", ""),
        ("Project URL", "https://_____.supabase.co"),
        ("Region", ""),
        ("Database password", "[PASSWORD MANAGER]"),
        ("Anon key (public)", "eyJ..."),
        ("Service role key (SECRET)", "eyJ... — never in frontend"),
        ("PostGIS enabled", "YES / NO"),
        ("Storage bucket", "inspection-files"),
        ("Edge functions deployed", "YES / NO"),
        ("CRON_SECRET", "[PASSWORD MANAGER]"),
        ("WEBHOOK_SECRET", "[PASSWORD MANAGER]"),
    ])

    doc.add_heading("Vercel (Web dashboard)", 1)
    add_table_rows(doc, ["Field", "Value"], [
        ("Team / account", ""),
        ("Project name", ""),
        ("Production URL", ""),
        ("Custom domain", ""),
        ("VERCEL_ORG_ID", ""),
        ("VERCEL_PROJECT_ID", ""),
        ("VERCEL_TOKEN", "[PASSWORD MANAGER]"),
        ("VITE_SUPABASE_URL", ""),
        ("VITE_SUPABASE_ANON_KEY", "eyJ..."),
    ])

    doc.add_heading("Expo / EAS (Mobile)", 1)
    add_table_rows(doc, ["Field", "Value"], [
        ("Expo account (owner)", ""),
        ("EAS project ID", ""),
        ("EXPO_TOKEN", "[PASSWORD MANAGER]"),
        ("Android package", "com.vigilance.kerala"),
        ("Production APK build URL", ""),
        ("OTA update channel", "production"),
    ])

    doc.add_heading("Resend (Email alerts)", 1)
    add_table_rows(doc, ["Field", "Value"], [
        ("Account email", ""),
        ("API key", "re_... [PASSWORD MANAGER]"),
        ("Verified sender (RESEND_FROM)", "VMS Alerts <...>"),
        ("Domain verified", "YES / NO"),
    ])

    doc.add_heading("Upstash (Rate limiting — optional)", 1)
    add_table_rows(doc, ["Field", "Value"], [
        ("Redis database name", ""),
        ("UPSTASH_REDIS_REST_URL", ""),
        ("UPSTASH_REDIS_REST_TOKEN", "[PASSWORD MANAGER]"),
    ])

    doc.add_heading("Old personal accounts — decommission after verification", 1)
    add_table_rows(doc, ["Service", "Old account", "Decommission date", "Status"], [
        ("Supabase", "itxfffjepcdfhuzsrnwf (Dhanushmcmsd)", "", "PENDING"),
        ("Vercel", "dhanushmcmsd / vigilance-web", "", "PENDING"),
        ("Expo", "dhanushraghav", "", "PENDING"),
        ("GitHub", "Dhanushmcmsd/Vigilance", "", "PENDING"),
        ("Resend", "", "", "PENDING"),
    ])

    doc.add_heading("Smoke test results (fill after cutover)", 1)
    tests = [
        "Admin web login",
        "Officer mobile login",
        "Inspection submit + geofence",
        "Photo upload to storage",
        "Red alert email",
        "Weekly report / cron",
        "CSV export",
        "Realtime dashboard updates",
    ]
    add_table_rows(doc, ["Test", "Pass/Fail", "Date", "Notes"], [(t, "", "", "") for t in tests])

    doc.add_heading("Transfer completion sign-off", 1)
    doc.add_paragraph("Transferred by: _________________________  Date: __________")
    doc.add_paragraph("Accepted by (company): ___________________  Date: __________")
    doc.add_paragraph()
    p = doc.add_paragraph("— End of confidential document —")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    build()
