#!/usr/bin/env python3
"""Generate Vigilance VMS handoff Word document."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import date

OUTPUT = r"e:\vigilance git\Vigilance\docs\Vigilance_Handoff_Transfer_Guide.docx"


def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    shading = p._element.get_or_add_pPr()
    return p


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def build():
    doc = Document()

    # Title
    title = doc.add_heading("Vigilance Management System (VMS)", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Complete Handoff & Transfer Guide")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.size = Pt(14)
    meta = doc.add_paragraph(f"Prepared: {date.today().strftime('%B %d, %Y')}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_paragraph(
        "This document describes how to transfer the Vigilance app — source code, "
        "Supabase database (schema + data + auth + storage), and Vercel web deployment — "
        "from a personal developer account to a company-owned account with no personal "
        "accounts linked."
    )

    # TOC-style sections
    add_heading(doc, "Table of Contents", 1)
    for item in [
        "1. Overview & What Gets Transferred",
        "2. Prerequisites (Accounts the Company Must Create)",
        "3. Phase 1 — Transfer Source Code (GitHub)",
        "4. Phase 2 — Transfer Supabase Database",
        "5. Phase 3 — Transfer Storage Files",
        "6. Phase 4 — Deploy Edge Functions & Secrets",
        "7. Phase 5 — Transfer Vercel Web Deployment",
        "8. Phase 6 — Mobile App (Expo EAS) — Brief",
        "9. Environment Variables Master Checklist",
        "10. Cutover & Verification",
        "11. Decommission Personal Accounts",
        "Appendix A — Storage Migration Script",
        "Appendix B — Files to Update in Codebase",
    ]:
        doc.add_paragraph(item, style="List Number")

    # 1. Overview
    add_heading(doc, "1. Overview & What Gets Transferred", 1)
    doc.add_paragraph("The VMS stack consists of:")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Component"
    hdr[1].text = "Technology"
    hdr[2].text = "Transfer Method"
    rows = [
        ("Source code", "GitHub monorepo (web/, mobile/, supabase/)", "Repo transfer or mirror"),
        ("Database", "Supabase PostgreSQL + PostGIS", "pg_dump / restore"),
        ("Auth users", "Supabase Auth (auth schema)", "pg_dump auth tables"),
        ("File storage", "Supabase bucket: inspection-files", "Scripted copy (Appendix A)"),
        ("Backend logic", "8 Supabase Edge Functions", "supabase functions deploy"),
        ("Web dashboard", "Vite + React on Vercel", "New Vercel project + env vars"),
        ("Mobile app", "Expo / EAS", "New EAS project + rebuild APK"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = val

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Recommended approach: ")
    r.bold = True
    p.add_run(
        "Create fresh company-owned accounts (Option B). Export data from your "
        "personal Supabase project, deploy to the company's new project, and point "
        "Vercel + mobile at the new credentials. This fully decouples your personal accounts."
    )

    # 2. Prerequisites
    add_heading(doc, "2. Prerequisites (Accounts the Company Must Create)", 1)
    for item in [
        "GitHub Organization (or company GitHub account)",
        "Supabase Organization + new Project (enable PostGIS extension)",
        "Vercel Team account",
        "Expo account (for mobile builds — optional at cutover if only web is needed first)",
        "Resend account (for alert/report emails)",
        "Upstash Redis account (for rate limiting — optional, degrades gracefully if missing)",
    ]:
        add_bullet(doc, item)

    doc.add_paragraph("Tools required on the migration machine:")
    for item in [
        "Node.js 20+, npm 10+",
        "Supabase CLI: npm install -g supabase",
        "PostgreSQL client tools: pg_dump and psql (from PostgreSQL installer or pgAdmin)",
        "Vercel CLI: npm install -g vercel",
        "Git",
        "Python 3 (for storage migration script in Appendix A)",
    ]:
        add_bullet(doc, item)

    # 3. GitHub
    add_heading(doc, "3. Phase 1 — Transfer Source Code (GitHub)", 1)

    add_heading(doc, "Method A — Transfer existing repository (keeps history)", 2)
    doc.add_paragraph("Performed by the current repo owner:")
    add_code_block(doc, """# On GitHub.com:
# 1. Open your repository → Settings → General
# 2. Scroll to "Danger Zone" → Transfer ownership
# 3. Enter the company GitHub org name and confirm
# 4. Company admin accepts the transfer invitation""")

    add_heading(doc, "Method B — Mirror to a fresh company repo (clean history optional)", 2)
    add_code_block(doc, """# On your machine — create a bare mirror and push to company repo
git clone --mirror https://github.com/YOUR_USER/Vigilance.git vigilance-mirror
cd vigilance-mirror
git remote set-url origin https://github.com/COMPANY_ORG/Vigilance.git
git push --mirror

# Company then clones:
git clone https://github.com/COMPANY_ORG/Vigilance.git
cd Vigilance
npm install""")

    add_heading(doc, "Remove personal account references in code (before or right after transfer)", 2)
    add_bullet(doc, " — change default Expo owner to company account", bold_prefix="mobile/app.config.js")
    add_bullet(doc, " — update site_url and redirect URLs to company Vercel domain", bold_prefix="supabase/config.toml")
    add_bullet(doc, " — ensure no .env files with secrets are committed (already in .gitignore)", bold_prefix=".env files")

    add_heading(doc, "GitHub Actions secrets (company sets in repo Settings → Secrets)", 2)
    for s in [
        "VERCEL_TOKEN — from vercel.com/account/tokens",
        "VERCEL_ORG_ID — from .vercel/project.json after linking",
        "VERCEL_PROJECT_ID — from .vercel/project.json after linking",
        "VITE_SUPABASE_URL — new Supabase project URL",
        "VITE_SUPABASE_ANON_KEY — new Supabase anon key",
        "EXPO_TOKEN — from expo.dev/settings/access-tokens (if using mobile CI)",
    ]:
        add_bullet(doc, s)

    # 4. Supabase DB
    add_heading(doc, "4. Phase 2 — Transfer Supabase Database", 1)

    doc.add_paragraph(
        "Current personal project reference (for export only): itxfffjepcdfhuzsrnwf. "
        "The company creates a NEW project and receives a new project ref (e.g. abcdefghijklmnop)."
    )

    add_heading(doc, "Step 4.1 — Company creates new Supabase project", 2)
    add_code_block(doc, """# Company account:
# 1. Go to https://supabase.com/dashboard → New Project
# 2. Choose region closest to users (e.g. South Asia if applicable)
# 3. Set a strong database password — save it securely
# 4. After creation: Database → Extensions → enable "postgis"
# 5. Note down from Settings → API:
#    - Project URL  (https://NEW_REF.supabase.co)
#    - anon public key
#    - service_role key (NEVER put in frontend — server/edge only)""")

    add_heading(doc, "Step 4.2 — Deploy schema from repository (migrations)", 2)
    add_code_block(doc, """# On a machine with the repo cloned:
cd Vigilance

# Login with COMPANY Supabase account
supabase login

# Link to the NEW company project (one-time)
supabase link --project-ref NEW_REF

# Apply all migrations in supabase/migrations/ to the new project
supabase db push

# Verify: Dashboard → Table Editor — tables like inspections, branches, stores should exist""")

    add_heading(doc, "Step 4.3 — Export data from OLD project (your account)", 2)
    doc.add_paragraph(
        "Get the database connection string from OLD project: "
        "Dashboard → Settings → Database → Connection string → URI (use direct connection, port 5432)."
    )
    add_code_block(doc, """# Replace placeholders:
# OLD_REF     = your current project ref (itxfffjepcdfhuzsrnwf)
# OLD_PASSWORD = database password from old project

# Export PUBLIC schema data only (tables, not structure — structure already applied via migrations)
pg_dump "postgresql://postgres.OLD_REF:OLD_PASSWORD@aws-0-ap-south-1.pooler.supabase.com:5432/postgres" ^
  --schema=public ^
  --data-only ^
  --no-owner ^
  --disable-triggers ^
  -f public_data_export.sql

# On macOS/Linux use backslashes instead of ^ for line continuation:
# pg_dump "postgresql://postgres.OLD_REF:OLD_PASSWORD@db.OLD_REF.supabase.co:5432/postgres" \\
#   --schema=public --data-only --no-owner --disable-triggers -f public_data_export.sql""")

    doc.add_paragraph(
        "Note: Use the direct connection host db.OLD_REF.supabase.co if pooler fails for pg_dump. "
        "Find the exact host in Supabase Dashboard → Settings → Database."
    )

    add_heading(doc, "Step 4.4 — Export AUTH users (so logins keep working)", 2)
    doc.add_paragraph(
        "This preserves password hashes. All active sessions will still be invalidated "
        "because the new project has a different JWT secret — users must log in again once."
    )
    add_code_block(doc, """# Export auth schema data from OLD project
pg_dump "postgresql://postgres.OLD_REF:OLD_PASSWORD@db.OLD_REF.supabase.co:5432/postgres" ^
  --schema=auth ^
  --data-only ^
  --no-owner ^
  -f auth_data_export.sql

# Tables included: auth.users, auth.identities, auth.sessions (sessions optional — can skip)""")

    add_heading(doc, "Step 4.5 — Import data into NEW company project", 2)
    add_code_block(doc, """# Replace NEW_REF and NEW_PASSWORD with company project values

# Import public data
psql "postgresql://postgres.NEW_REF:NEW_PASSWORD@db.NEW_REF.supabase.co:5432/postgres" ^
  -f public_data_export.sql

# Import auth users
psql "postgresql://postgres.NEW_REF:NEW_PASSWORD@db.NEW_REF.supabase.co:5432/postgres" ^
  -f auth_data_export.sql

# If you get foreign-key order errors, re-run after public import completes.
# Alternative: combine both dumps or use --disable-triggers on import.""")

    add_heading(doc, "Step 4.6 — Alternative: Full database backup/restore (advanced)", 2)
    add_code_block(doc, """# Full custom-format dump (schema + data) from old project:
pg_dump "postgresql://postgres.OLD_REF:OLD_PASSWORD@db.OLD_REF.supabase.co:5432/postgres" ^
  --format=custom ^
  --no-owner ^
  --schema=public ^
  --schema=auth ^
  -f vigilance_full.backup

# Restore to new project (run AFTER supabase db push so extensions exist):
pg_restore --dbname="postgresql://postgres.NEW_REF:NEW_PASSWORD@db.NEW_REF.supabase.co:5432/postgres" ^
  --no-owner ^
  --data-only ^
  vigilance_full.backup

# Prefer Step 4.2 + 4.3 + 4.5 for cleaner migration with version-controlled migrations.""")

    add_heading(doc, "Step 4.7 — Update Supabase Auth redirect URLs", 2)
    add_code_block(doc, """# Edit supabase/config.toml in the repo, then push config:
# Change site_url to company Vercel URL, e.g.:
#   site_url = "https://vigilance.company.com"
# Update additional_redirect_urls with company domains

supabase link --project-ref NEW_REF
supabase config push

# Or set manually in Dashboard → Authentication → URL Configuration""")

    # 5. Storage
    add_heading(doc, "5. Phase 3 — Transfer Storage Files", 1)
    doc.add_paragraph(
        "Inspection photos and attachments live in the private Supabase Storage bucket "
        "named inspection-files. SQL dumps do NOT include these files."
    )
    add_code_block(doc, """# Quick manual check — list files in old project (Dashboard → Storage → inspection-files)
# For large volumes, use the Python script in Appendix A.

# Prerequisites for Appendix A script:
pip install supabase

# Set environment variables then run:
# OLD: SUPABASE_URL + SUPABASE_SERVICE_ROLE_KEY (old project)
# NEW: NEW_SUPABASE_URL + NEW_SUPABASE_SERVICE_ROLE_KEY (new project)
python scripts/migrate-storage.py""")

    # 6. Edge functions
    add_heading(doc, "6. Phase 4 — Deploy Edge Functions & Secrets", 1)

    add_heading(doc, "Edge functions in this project", 2)
    for fn in [
        "admin-create-user", "admin-update-user", "export-csv", "health-check",
        "notify-officer", "on-inspection-submit", "red-alert", "weekly-report",
    ]:
        add_bullet(doc, fn)

    add_code_block(doc, """# Deploy all edge functions to company project
cd Vigilance
supabase link --project-ref NEW_REF
supabase functions deploy

# Or deploy one at a time:
supabase functions deploy on-inspection-submit
supabase functions deploy red-alert
supabase functions deploy weekly-report
# ... etc.""")

    add_heading(doc, "Set all Supabase Edge secrets (company account)", 2)
    add_code_block(doc, """supabase secrets set ^
  SUPABASE_URL=https://NEW_REF.supabase.co ^
  SUPABASE_ANON_KEY=eyJ...company_anon_key... ^
  SUPABASE_SERVICE_ROLE_KEY=eyJ...company_service_role... ^
  RESEND_API_KEY=re_...company_resend_key... ^
  RESEND_FROM="VMS Alerts <alerts@company-domain.com>" ^
  DASHBOARD_URL=https://company-web-url.vercel.app ^
  UPSTASH_REDIS_REST_URL=https://....upstash.io ^
  UPSTASH_REDIS_REST_TOKEN=... ^
  CRON_SECRET=generate_a_long_random_string_here ^
  WEBHOOK_SECRET=generate_another_random_string_here

# Verify secrets (names only, not values):
supabase secrets list""")

    doc.add_paragraph(
        "Re-create any Database Webhooks in the new Supabase Dashboard that trigger "
        "on-inspection-submit (if configured on the old project). "
        "Dashboard → Database → Webhooks."
    )

    # 7. Vercel
    add_heading(doc, "7. Phase 5 — Transfer Vercel Web Deployment", 1)

    add_heading(doc, "Method A — New Vercel project under company account (recommended)", 2)
    add_code_block(doc, """# Company team member:
npm install -g vercel
vercel login

cd Vigilance/web

# Link to new Vercel project (interactive — choose company team)
vercel link

# Set production environment variables
vercel env add VITE_SUPABASE_URL production
# Paste: https://NEW_REF.supabase.co

vercel env add VITE_SUPABASE_ANON_KEY production
# Paste: company anon key (eyJ...)

# Pull env locally to verify
vercel env pull .env.production.local

# Deploy to production
vercel --prod

# Note the deployed URL, e.g. https://vigilance-web-company.vercel.app
# Update DASHBOARD_URL in Supabase secrets to match.""")

    add_heading(doc, "Method B — Transfer existing Vercel project between accounts", 2)
    doc.add_paragraph(
        "Vercel supports transferring a project to another team if you are the owner:"
    )
    add_code_block(doc, """# On vercel.com:
# 1. Open project → Settings → General
# 2. Scroll to "Transfer Project"
# 3. Select the company Vercel team
# 4. Company admin accepts

# After transfer, update environment variables in the NEW team's project settings:
#   VITE_SUPABASE_URL → new Supabase URL
#   VITE_SUPABASE_ANON_KEY → new anon key
# Then redeploy:
vercel --prod""")

    add_heading(doc, "Method C — CI/CD via GitHub Actions (already in repo)", 2)
    add_code_block(doc, """# File: .github/workflows/deploy-web.yml
# Company adds these GitHub Secrets, then push to main triggers deploy:

VERCEL_TOKEN=...          # vercel.com/account/tokens
VERCEL_ORG_ID=...         # from web/.vercel/project.json → orgId
VERCEL_PROJECT_ID=...     # from web/.vercel/project.json → projectId
VITE_SUPABASE_URL=https://NEW_REF.supabase.co
VITE_SUPABASE_ANON_KEY=eyJ...

# Get org/project IDs after vercel link:
cat web/.vercel/project.json""")

    add_heading(doc, "Custom domain (optional)", 2)
    add_code_block(doc, """# Vercel Dashboard → Project → Settings → Domains → Add
# e.g. vigilance.company.com
# Add DNS CNAME record pointing to cname.vercel-dns.com
# Update supabase/config.toml site_url and redirect URLs to match
# Update DASHBOARD_URL secret in Supabase""")

    # 8. Mobile brief
    add_heading(doc, "8. Phase 6 — Mobile App (Expo EAS) — Brief", 1)
    add_code_block(doc, """cd mobile
eas login                    # company Expo account
eas init                     # creates new EAS project ID

# Set in mobile/.env (or EAS secrets):
EXPO_PUBLIC_SUPABASE_URL=https://NEW_REF.supabase.co
EXPO_PUBLIC_SUPABASE_ANON_KEY=eyJ...

# Update mobile/app.config.js defaults:
#   expoOwner → company Expo username
#   easProjectId → new project ID from eas init

eas build --platform android --profile production
# Distribute new APK to field officers — old APK still points at old Supabase until replaced""")

    # 9. Env checklist
    add_heading(doc, "9. Environment Variables Master Checklist", 1)
    env_table = doc.add_table(rows=1, cols=4)
    env_table.style = "Table Grid"
    eh = env_table.rows[0].cells
    eh[0].text = "Variable"
    eh[1].text = "Where to set"
    eh[2].text = "Used by"
    eh[3].text = "Notes"
    env_rows = [
        ("VITE_SUPABASE_URL", "Vercel + web/.env", "Web dashboard", "Public — safe in browser"),
        ("VITE_SUPABASE_ANON_KEY", "Vercel + web/.env", "Web dashboard", "Public — RLS protects data"),
        ("EXPO_PUBLIC_SUPABASE_URL", "mobile/.env + EAS", "Mobile app", "Baked into APK at build"),
        ("EXPO_PUBLIC_SUPABASE_ANON_KEY", "mobile/.env + EAS", "Mobile app", "Baked into APK at build"),
        ("SUPABASE_SERVICE_ROLE_KEY", "Supabase Edge secrets ONLY", "Edge functions", "NEVER in frontend"),
        ("RESEND_API_KEY", "Supabase Edge secrets", "Email functions", "Company Resend account"),
        ("RESEND_FROM", "Supabase Edge secrets", "Email functions", "Must be verified domain"),
        ("DASHBOARD_URL", "Supabase Edge secrets", "Email links", "Company Vercel URL"),
        ("UPSTASH_REDIS_REST_URL", "Supabase Edge secrets", "Rate limiting", "Optional"),
        ("UPSTASH_REDIS_REST_TOKEN", "Supabase Edge secrets", "Rate limiting", "Optional"),
        ("CRON_SECRET", "Supabase Edge secrets", "health-check, cron", "Random string"),
        ("WEBHOOK_SECRET", "Supabase Edge secrets", "on-inspection-submit", "Match DB webhook header"),
        ("VERCEL_TOKEN", "GitHub Actions secrets", "CI deploy", "Company Vercel token"),
        ("VERCEL_ORG_ID", "GitHub Actions secrets", "CI deploy", "From .vercel/project.json"),
        ("VERCEL_PROJECT_ID", "GitHub Actions secrets", "CI deploy", "From .vercel/project.json"),
        ("EXPO_TOKEN", "GitHub Actions secrets", "Mobile CI", "Company Expo token"),
    ]
    for row in env_rows:
        cells = env_table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val

    # 10. Cutover
    add_heading(doc, "10. Cutover & Verification", 1)
    doc.add_paragraph("Recommended maintenance window: 1–2 hours.")
    for i, step in enumerate([
        "Freeze writes on old system (optional: set old Supabase to read-only or announce downtime).",
        "Run final pg_dump of public + auth data from old project.",
        "Import into new company Supabase project.",
        "Run storage migration script (Appendix A).",
        "Deploy edge functions + verify secrets.",
        "Deploy Vercel with new env vars.",
        "Update Supabase Auth redirect URLs to company Vercel domain.",
        "Build and distribute new mobile APK (if mobile is in use).",
        "Smoke test all roles: admin login, officer inspection submit, file upload, email alert.",
        "Monitor for 7 days, then decommission personal accounts.",
    ], 1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")

    add_heading(doc, "Smoke test checklist", 2)
    for t in [
        "Admin can log in on web dashboard",
        "Field officer can log in on mobile",
        "Geofenced inspection submission works",
        "Photo upload to inspection-files bucket works",
        "Red alert email sends (test with a high-risk inspection)",
        "Weekly report cron/edge function runs",
        "CSV export works for management role",
        "Realtime updates appear on dashboard",
    ]:
        add_bullet(doc, t)

    # 11. Decommission
    add_heading(doc, "11. Decommission Personal Accounts", 1)
    decom = doc.add_table(rows=1, cols=2)
    decom.style = "Table Grid"
    decom.rows[0].cells[0].text = "Account"
    decom.rows[0].cells[1].text = "Action after verified cutover"
    decom_data = [
        ("Supabase (old project)", "Pause then delete project in Dashboard"),
        ("Vercel (old project)", "Delete project or complete transfer"),
        ("Expo (old EAS project)", "Archive or leave inactive"),
        ("Resend", "Revoke API key, remove domain"),
        ("Upstash", "Delete Redis database"),
        ("GitHub", "Transfer complete — you lose push access"),
    ]
    for a, b in decom_data:
        cells = decom.add_row().cells
        cells[0].text = a
        cells[1].text = b

    # Appendix A
    add_heading(doc, "Appendix A — Storage Migration Script", 1)
    doc.add_paragraph("Save as scripts/migrate-storage.py and run with company + old credentials:")
    add_code_block(doc, '''#!/usr/bin/env python3
"""Copy all files from old Supabase storage bucket to new project."""
import os
from supabase import create_client

OLD_URL = os.environ["OLD_SUPABASE_URL"]
OLD_KEY = os.environ["OLD_SUPABASE_SERVICE_ROLE_KEY"]
NEW_URL = os.environ["NEW_SUPABASE_URL"]
NEW_KEY = os.environ["NEW_SUPABASE_SERVICE_ROLE_KEY"]
BUCKET = "inspection-files"

old_sb = create_client(OLD_URL, OLD_KEY)
new_sb = create_client(NEW_URL, NEW_KEY)

# Ensure bucket exists on new project (migrations should create it)
try:
    new_sb.storage.create_bucket(BUCKET, options={"public": False})
except Exception:
    pass  # already exists

# List all files recursively
def list_all(prefix=""):
    items = old_sb.storage.from_(BUCKET).list(prefix)
    paths = []
    for item in items:
        path = f"{prefix}/{item['name']}".lstrip("/")
        if item.get("id"):  # file
            paths.append(path)
        else:  # folder
            paths.extend(list_all(path))
    return paths

files = list_all()
print(f"Found {len(files)} files to migrate")

for i, path in enumerate(files, 1):
  try:
    data = old_sb.storage.from_(BUCKET).download(path)
    new_sb.storage.from_(BUCKET).upload(path, data, {"upsert": "true"})
    print(f"[{i}/{len(files)}] OK: {path}")
  except Exception as e:
    print(f"[{i}/{len(files)}] FAIL: {path} — {e}")

print("Done.")''')

    add_code_block(doc, """# Run:
set OLD_SUPABASE_URL=https://itxfffjepcdfhuzsrnwf.supabase.co
set OLD_SUPABASE_SERVICE_ROLE_KEY=your_old_service_role_key
set NEW_SUPABASE_URL=https://NEW_REF.supabase.co
set NEW_SUPABASE_SERVICE_ROLE_KEY=company_service_role_key
pip install supabase
python scripts/migrate-storage.py""")

    # Appendix B
    add_heading(doc, "Appendix B — Files to Update in Codebase", 1)
    files_table = doc.add_table(rows=1, cols=2)
    files_table.style = "Table Grid"
    files_table.rows[0].cells[0].text = "File"
    files_table.rows[0].cells[1].text = "What to change"
    file_data = [
        ("mobile/app.config.js", "expoOwner, easProjectId defaults → company values"),
        ("supabase/config.toml", "site_url, additional_redirect_urls → company domains"),
        ("web/.env.example", "Documentation only — company copies to .env"),
        ("mobile/.env.example", "Documentation only — company copies to .env"),
        (".github/workflows/*.yml", "Secrets configured in GitHub, not in files"),
    ]
    for f, c in file_data:
        cells = files_table.add_row().cells
        cells[0].text = f
        cells[1].text = c

  # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph(
        "— End of document — Vigilance Management System Handoff Guide"
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    build()
