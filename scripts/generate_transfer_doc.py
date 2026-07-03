#!/usr/bin/env python3
"""Generate VMS_Company_Transfer_Complete.docx with actual transfer results."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import date
from pathlib import Path

OUTPUT = r"e:\vigilance git\Vigilance\VMS_Company_Transfer_Complete.docx"

NAVY = RGBColor(0x1A, 0x23, 0x7E)
GREEN = RGBColor(0x00, 0x69, 0x5C)
RED = RGBColor(0xC6, 0x28, 0x28)
ALT = "E8EAF6"
BORDER = "9FA8DA"


def shade_cell(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def style_header_row(table):
    for cell in table.rows[0].cells:
        shade_cell(cell, "1A237E")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.bold = True


def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    style_header_row(t)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            if ri % 2 == 1:
                shade_cell(cells[ci], ALT)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def add_credential_card(doc, email, password, role):
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    left, right = t.rows[0].cells
    shade_cell(left, "E8EAF6")
    left.paragraphs[0].add_run(f"{email}\n{password}").bold = True
    p = right.paragraphs[0]
    r = p.add_run(role.upper())
    r.bold = True
    r.font.color.rgb = NAVY
    doc.add_paragraph()


def build():
    doc = Document()

    # Cover
    cover = doc.add_table(rows=1, cols=1)
    cell = cover.rows[0].cells[0]
    shade_cell(cell, "1A237E")
    for line, size, bold in [
        ("VIGILANCE MANAGEMENT SYSTEM", 22, True),
        ("Company Account Transfer — Completion Report", 14, True),
        ("Version 1.1.0  |  June 20, 2026", 11, False),
        ("Prepared for: CFC ICI Kerala", 11, False),
    ]:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(size)
        r.bold = bold
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_page_break()

    doc.add_heading("Section 1 — Transfer Summary", level=1)
    add_table(
        doc,
        ["Item", "Old (Developer)", "New (Company)"],
        [
            ("GitHub Repo", "Dhanushmcmsd/Vigilance", "vigilancecfcici/VigilanceManagementSystem"),
            ("Supabase Project Ref", "itxfffjepcdfhuzsrnwf", "xgwjcknpkpzsbjvuninm"),
            ("Supabase URL", "https://itxfffjepcdfhuzsrnwf.supabase.co", "https://xgwjcknpkpzsbjvuninm.supabase.co"),
            ("Vercel URL", "vigilance-web.vercel.app (personal)", "https://web-two-orcin-uy22un2bzf.vercel.app"),
            ("Expo Account", "dhanushraghav", "vigilancecfcici"),
            ("EAS Project ID", "b63bc80d-778c-4eb8-b0f7-b974238318e6", "aa8ccd7a-d071-4d7e-9b44-ecc1192f7b06"),
            ("Transfer Date", "—", "June 20, 2026"),
            ("Total Commits", "—", "69"),
            ("DB Migrations", "55", "55"),
            ("Edge Functions", "9", "9 (all ACTIVE)"),
        ],
    )

    doc.add_heading("Section 2 — Phase-by-Phase Results", level=1)
    add_table(
        doc,
        ["Phase", "Description", "Status", "Notes"],
        [
            ("1", "GitHub code push", "✅", "69 commits on main"),
            ("2", "Supabase schema bootstrap", "✅", "pg_dump baseline + 55 migrations marked applied"),
            ("3", "Data + auth transfer", "✅", "80 stores, 91 branches, 5 users"),
            ("4", "Edge functions (9)", "✅", "Including supervisor-otp restored from old project"),
            ("5", "Secrets / env vars", "✅", "CRON_SECRET + WEBHOOK_SECRET generated; EAS + Vercel dev env added"),
            ("6", "Vercel deployment", "✅", "HTTP 200 on production alias"),
            ("7", "Expo EAS setup", "✅", "Build 9e86637f-83fa-4f41-9aa1-10c24890cde2 queued"),
            ("8", "Final verification", "✅", "All critical checks passed"),
        ],
    )

    doc.add_heading("Section 3 — Database Verification", level=1)
    add_table(
        doc,
        ["Table / Check", "Expected", "Actual", "Status"],
        [
            ("public tables", "20 app tables (+ spatial_ref_sys in extensions)", "20", "✅"),
            ("stores", "80", "80", "✅"),
            ("branches", "91", "91", "✅"),
            ("checklist_templates", "32", "32", "✅"),
            ("user_roles", "5", "5", "✅"),
            ("auth.users", "5", "5", "✅"),
            ("inspections", "0", "0", "✅"),
            ("migrations", "55", "55", "✅"),
            ("postgis", "1", "1", "✅"),
            ("was_officer_in_range()", "1", "1", "✅"),
            ("location_status column", "1", "1", "✅"),
        ],
    )

    doc.add_heading("Section 4 — Credential Cards", level=1)
    doc.add_paragraph("Change all passwords immediately after first login.")
    for email, pw, role in [
        ("admin@vigilance.app", "Admin@2026", "admin"),
        ("management@vigilance.app", "Mgmt@2026", "management"),
        ("officer@vigilance.app", "Officer@2026", "officer"),
        ("audit@company.app", "Audit@2026", "audit"),
    ]:
        add_credential_card(doc, email, pw, role)
    doc.add_paragraph("Additional transferred user: niyas.ts@cfcici.co.in (officer)")

    doc.add_heading("Section 5 — API Reference", level=1)
    add_table(
        doc,
        ["Key", "Value"],
        [
            ("Supabase Project URL", "https://xgwjcknpkpzsbjvuninm.supabase.co"),
            ("Supabase Anon Key", "sb_publishable_X0hpkdA…zakhiDdF (truncated)"),
            ("Supabase Dashboard", "https://supabase.com/dashboard/project/xgwjcknpkpzsbjvuninm"),
            ("GitHub Repo", "https://github.com/vigilancecfcici/VigilanceManagementSystem"),
            ("Vercel Production URL", "https://web-two-orcin-uy22un2bzf.vercel.app"),
            ("Vercel Dashboard", "https://vercel.com/vigilancecfcici1/web"),
            ("Expo Dashboard", "https://expo.dev/accounts/vigilancecfcici/projects/vigilance-management-system"),
            ("EAS Build", "https://expo.dev/accounts/vigilancecfcici/projects/vigilance-management-system/builds/9e86637f-83fa-4f41-9aa1-10c24890cde2"),
        ],
    )
    p = doc.add_paragraph()
    r = p.add_run("Rotate all API tokens and developer credentials now. Only the anon/publishable key belongs in web/mobile config.")
    r.font.color.rgb = RED
    r.bold = True

    doc.add_heading("Section 6 — Geofencing Status", level=1)
    for line in [
        "PostGIS: ✅ Active",
        "was_officer_in_range(): ✅ Deployed",
        "compute_inspection_location_status(): ✅ Deployed",
        "location_status column: ✅ Present on inspections",
        "on-inspection-submit: ✅ Deployed (geofence + location_status write)",
        "Server-side enforcement: ✅ Active for new submissions",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Section 7 — Post-Transfer Security Checklist", level=1)
    for item in [
        "Rotate GitHub PAT used for transfer",
        "Rotate Supabase access token and service role key",
        "Rotate Vercel token",
        "Rotate Expo token",
        "Set RESEND_API_KEY, UPSTASH_REDIS_REST_URL, UPSTASH_REDIS_REST_TOKEN on new Supabase project",
        "Change all 5 user passwords before handing to company staff",
        "Pause or delete old Supabase project itxfffjepcdfhuzsrnwf after verification period",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Section 8 — Resend / Email Function Audit", level=1)
    add_table(
        doc,
        ["Function", "RESEND ref?", "Vault secret?", "Decision"],
        [
            ("notify-officer", "Yes", "No", "Deploy — graceful skip without key"),
            ("weekly-report", "Yes", "No", "Deploy — needs RESEND_API_KEY for emails"),
            ("red-alert", "Yes", "No", "Deploy — needs RESEND_API_KEY for emails"),
            ("on-inspection-submit", "Yes", "No", "Deploy — geofence works; emails need key"),
            ("supervisor-otp", "Yes", "No", "Deploy — OTP email needs key"),
        ],
    )

    doc.add_heading("Appendix A — Transferred Auth Users", level=1)
    add_table(
        doc,
        ["Email", "Role"],
        [
            ("admin@vigilance.app", "admin"),
            ("audit@company.app", "audit"),
            ("management@vigilance.app", "management"),
            ("niyas.ts@cfcici.co.in", "officer"),
            ("officer@vigilance.app", "officer"),
        ],
    )

    doc.add_heading("Appendix B — Deployed Edge Functions", level=1)
    add_table(
        doc,
        ["Function", "Status", "JWT"],
        [
            ("admin-create-user", "ACTIVE", "required"),
            ("admin-update-user", "ACTIVE", "required"),
            ("export-csv", "ACTIVE", "required"),
            ("health-check", "ACTIVE", "disabled"),
            ("notify-officer", "ACTIVE", "required"),
            ("on-inspection-submit", "ACTIVE", "required"),
            ("red-alert", "ACTIVE", "required"),
            ("supervisor-otp", "ACTIVE", "required"),
            ("weekly-report", "ACTIVE", "required"),
        ],
    )

    doc.add_heading("Appendix C — Data Tables Transferred", level=1)
    add_table(
        doc,
        ["Table", "Rows"],
        [
            ("branch_types", "2"),
            ("stores", "80"),
            ("branches", "91"),
            ("risk_classifications", "32"),
            ("checklist_templates", "32"),
            ("user_roles", "5"),
            ("district_assignments", "12"),
            ("store_officer_assignments", "12"),
        ],
    )

    doc.add_paragraph(
        "Note: spatial_ref_sys (PostGIS reference data) lives in the extensions schema on the new "
        "Supabase project instead of public — this is expected platform behaviour and does not affect geofencing."
    )

    doc.add_heading("Section 9 — Environment Variables — Final State", level=1)

    doc.add_heading("9.1 Supabase Secrets (project xgwjcknpkpzsbjvuninm)", level=2)
    add_table(
        doc,
        ["Secret name", "Status", "Notes"],
        [
            ("CRON_SECRET", "✅ Set", "Stored in Supabase vault only — generated June 20, 2026"),
            ("WEBHOOK_SECRET", "✅ Set", "Stored in Supabase vault only — generated June 20, 2026"),
            ("DASHBOARD_URL", "✅ Set", "https://web-two-orcin-uy22un2bzf.vercel.app"),
            ("RESEND_FROM", "✅ Set", "VMS Alerts sender address"),
            ("SUPABASE_URL", "✅ Set", "Auto-managed by Supabase"),
            ("SUPABASE_ANON_KEY", "✅ Set", "Auto-managed by Supabase"),
            ("SUPABASE_SERVICE_ROLE_KEY", "✅ Set", "Auto-managed by Supabase"),
            ("RESEND_API_KEY", "⏳ PENDING", "Required from IT admin for email alerts"),
            ("UPSTASH_REDIS_REST_URL", "⏳ PENDING", "Optional — rate limiting degrades gracefully"),
            ("UPSTASH_REDIS_REST_TOKEN", "⏳ PENDING", "Optional — rate limiting degrades gracefully"),
        ],
    )

    doc.add_heading("9.2 Vercel Environment Variables (vigilancecfcici1/web)", level=2)
    add_table(
        doc,
        ["Variable", "Production", "Preview", "Development"],
        [
            ("VITE_SUPABASE_URL", "✅ [SENSITIVE]", "✅ [SENSITIVE]", "✅"),
            ("VITE_SUPABASE_ANON_KEY", "✅ [SENSITIVE]", "✅ [SENSITIVE]", "✅"),
        ],
    )
    doc.add_paragraph("CRON_SECRET and WEBHOOK_SECRET are not used by the web app — Supabase edge functions only.")

    doc.add_heading("9.3 Expo / EAS Environment Variables", level=2)
    add_table(
        doc,
        ["Variable", "Production", "Preview"],
        [
            ("EXPO_PUBLIC_SUPABASE_URL", "✅", "✅"),
            ("EXPO_PUBLIC_SUPABASE_ANON_KEY", "✅", "✅"),
        ],
    )

    doc.add_heading("9.4 Generated Secrets", level=2)
    doc.add_paragraph(
        "CRON_SECRET and WEBHOOK_SECRET were generated on June 20, 2026 using "
        "cryptographically secure random bytes (openssl rand -hex 32 equivalent). "
        "Values are stored in the Supabase secrets vault only and are NOT recorded in this document."
    )

    doc.add_heading("9.5 Pending / Still Needed from IT Admin", level=2)
    for item in [
        "RESEND_API_KEY — required for red-alert, weekly-report, on-inspection-submit emails, supervisor-otp",
        "UPSTASH_REDIS_REST_URL + UPSTASH_REDIS_REST_TOKEN — optional rate limiting",
        "Recreate Database Webhook on inspections → on-inspection-submit with x-webhook-secret header",
        "Migrate inspection-files storage bucket if historical photos are needed",
        "Rotate all transfer-session API tokens",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Appendix D — Remaining Manual Steps for IT Admin", level=1)
    steps = [
        "Provide RESEND_API_KEY and run: supabase secrets set RESEND_API_KEY=re_xxx --project-ref xgwjcknpkpzsbjvuninm",
        "Provide Upstash credentials for rate limiting (optional but recommended for production).",
        "Recreate Database Webhook on inspections → on-inspection-submit using WEBHOOK_SECRET as x-webhook-secret header.",
        "Migrate storage bucket inspection-files using scripts/migrate-storage.py if historical photos are required.",
        "Install the new Android APK from EAS build 9e86637f-83fa-4f41-9aa1-10c24890cde2.",
        "Rotate every token that was used during this transfer session.",
        "Change demo account passwords before distributing credentials to staff.",
    ]
    for s in steps:
        doc.add_paragraph(s, style="List Number")

    doc.add_heading("Appendix E — All 55 Migrations Applied", level=1)
    mig_rows = []
    mig_dir = Path(__file__).resolve().parents[1] / "supabase" / "migrations"
    for i, f in enumerate(sorted(mig_dir.glob("*.sql")), 1):
        mig_rows.append((str(i), f.name))
    add_table(doc, ["#", "Migration file"], mig_rows)

    doc.add_heading("Section 10 — Post-Transfer DB & App Health Verification", level=1)
    doc.add_paragraph("Verification date: June 20, 2026")

    doc.add_heading("10.1 Migrations", level=2)
    doc.add_paragraph(
        "Local migration files: 55. Database schema_migrations records: 55. "
        "Zero local-only or DB-only entries — full match."
    )

    doc.add_heading("10.2 Schema Drift", level=2)
    doc.add_paragraph(
        "TypeScript types regenerated from live DB (project xgwjcknpkpzsbjvuninm) and synced to "
        "web/src/types/supabase.ts and mobile/types/supabase.ts. "
        "Removed stale PostGIS view/function noise from old project types; "
        "app RPCs (was_officer_in_range, compute_inspection_location_status, etc.) now match live schema."
    )
    doc.add_paragraph(
        "Manual fix applied: was_officer_in_range and compute_inspection_location_status "
        "search_path set to public, extensions (PostGIS geometry type resolution)."
    )

    doc.add_heading("10.3 RLS Policies & Triggers", level=2)
    add_table(
        doc,
        ["Check", "Expected", "Actual", "Status"],
        [
            ("RLS policies (public)", "64 (old + new project)", "64", "✅"),
            ("Triggers (public)", "12", "12", "✅"),
        ],
    )
    doc.add_paragraph(
        "Note: Pre-flight report cited 73 policies; live count on both old and new projects is 64. "
        "All 20 public app tables have at least one policy — no tables with zero policies."
    )

    doc.add_heading("10.4 Geofence Functions", level=2)
    add_table(
        doc,
        ["Function", "Test", "Result"],
        [
            ("was_officer_in_range(uuid)", "Compiles and executes (auth guard active)", "✅"),
            ("compute_inspection_location_status(uuid)", "Present in pg_proc (1 arg)", "✅"),
            ("PostGIS ST_DWithin", "Same-point 100m radius", "✅ true"),
        ],
    )

    doc.add_heading("10.5 Edge Function Health", level=2)
    add_table(
        doc,
        ["Function", "HTTP Code", "Status", "Notes"],
        [
            ("health-check", "200", "✅", "GET + x-cron-secret"),
            ("on-inspection-submit", "200", "✅", "Skipped (test payload); not 500"),
            ("export-csv", "200", "✅", "GET + anon JWT"),
            ("supervisor-otp", "400", "✅", "Unknown action — function alive"),
            ("admin-create-user", "401", "✅", "Requires admin user JWT (not service role)"),
            ("admin-update-user", "401", "✅", "Requires admin user JWT (not service role)"),
        ],
    )
    doc.add_paragraph("Zero HTTP 500 responses across all tested functions.")

    doc.add_heading("10.6 Smoke Test Results", level=2)
    add_table(
        doc,
        ["Metric", "Expected", "Actual", "Status"],
        [
            ("public tables", "20", "20", "✅"),
            ("stores", "80", "80", "✅"),
            ("branches", "91", "91", "✅"),
            ("checklist_templates", "32", "32", "✅"),
            ("auth.users", "5", "5", "✅"),
            ("RLS policies", "64", "64", "✅"),
            ("triggers", "12", "12", "✅"),
            ("migrations", "55", "55", "✅"),
        ],
    )

    doc.add_heading("10.7 Vercel & EAS", level=2)
    add_table(
        doc,
        ["Service", "URL / ID", "Status"],
        [
            ("Vercel production", "https://web-two-orcin-uy22un2bzf.vercel.app", "HTTP 200 ✅"),
            (
                "EAS Android preview",
                "9e86637f-83fa-4f41-9aa1-10c24890cde2",
                "FINISHED ✅",
            ),
            (
                "APK download",
                "https://expo.dev/artifacts/eas/Q6RyDVFg-4LRXkD_VJR3b6hB7Ry5Nf_EGCn1jJKd0p0.apk",
                "Available",
            ),
        ],
    )
    doc.add_paragraph(
        "CRON_SECRET and WEBHOOK_SECRET were rotated during health verification (June 20, 2026). "
        "Values remain vault-only. Recreate the inspections database webhook with the new WEBHOOK_SECRET."
    )

    doc.add_heading("Section 11 — Support Contact", level=1)
    doc.add_paragraph("Developer: Dhanush")
    doc.add_paragraph("Original repo: github.com/Dhanushmcmsd/Vigilance")
    doc.add_paragraph("Handover date: June 20, 2026")

    doc.save(OUTPUT)
    size = __import__("os").path.getsize(OUTPUT)
    print(f"Created: {OUTPUT} ({size} bytes)")


if __name__ == "__main__":
    build()
